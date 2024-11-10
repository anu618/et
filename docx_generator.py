from docx import Document
from docx.shared import Pt, RGBColor
import re
import os
import logging
from docx.oxml import OxmlElement

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def clean_special_characters(text):
    """Clean up special characters and formatting"""
    if not text:
        return text
    
    # Remove URLs
    text = re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', '', text)
    
    # Remove timestamps and dates
    text = re.sub(r'\d{1,2}\s*(year|month|week|day)s?(,?\s*\d{1,2}\s*(year|month|week|day)s?)?\s*ago', '', text)
    
    # Remove voting information
    text = re.sub(r'(?i)upvoted\s*\d+\s*times?', '', text)
    text = re.sub(r'(?i)voted\s*\d+\s*times?', '', text)
    
    # Remove user comments and metadata
    text = re.sub(r'Selected Answer:.*', '', text)
    text = re.sub(r'Chosen Answer:.*', '', text)
    text = re.sub(r'Community.*', '', text)
    
    # Remove any remaining special characters and extra whitespace
    text = re.sub(r'[^\w\s.,():-]', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    
    return text.strip()

def parse_question(question_text):
    """Parse a single question section and return structured data"""
    try:
        # Extract question number
        number_match = re.search(r'Question #?:?\s*(\d+)', question_text)
        if not number_match:
            return None
        question_number = number_match.group(1)
        
        # Extract question content
        content_parts = question_text.split('--------------------------------------------------------------')
        if len(content_parts) < 3:
            return None
            
        # Clean up question content - take only the actual question
        question_content = content_parts[1].strip()
        
        # Extract options section
        options_text = content_parts[2].strip()
        
        # Clean up options - take only the actual options before any comments
        options = []
        for line in options_text.split('\n'):
            line = line.strip()
            if not line:
                continue
            # Only take lines that start with A., B., C., etc.
            if re.match(r'^[A-E]\.\s', line):
                # Clean up the option
                option = clean_special_characters(line)
                options.append(option)
        
        # Extract correct answer
        correct_match = re.search(r'CORRECT ANSWER==:\s*([A-D]+)', question_text)
        correct_answers = list(correct_match.group(1)) if correct_match else []
        
        return {
            'number': question_number,
            'content': question_content,
            'options': options,
            'correct': correct_answers
        }
        
    except Exception as e:
        print(f"Error parsing question: {str(e)}")
        return None

def format_cleaned_question(question_data):
    """Format a cleaned question for output"""
    lines = [
        f"Question #: {question_data['number']}",
        "--------------------------------------------------------------",
        question_data['content'].strip(),
        "--------------------------------------------------------------"
    ]
    
    # Add options
    for option in question_data['options']:
        lines.append(option.strip())
    
    # Add correct answer
    if question_data['correct']:
        lines.append("--------------------------------------------------------------")
        lines.append(f"CORRECT ANSWER==: {''.join(question_data['correct'])}")
    
    return '\n'.join(lines)

def create_practice_document(input_file, output_file):
    """Creates a practice document without answers"""
    try:
        doc = Document()
        
        # Read and parse input file
        with open(input_file, 'r', encoding='utf-8', errors='replace') as file:
            content = file.read()
            
        # Split into questions
        questions = content.split('###################################################################')
        
        for question_text in questions:
            if not question_text.strip():
                continue
                
            question_data = parse_question(question_text)
            if not question_data:
                continue
                
            # Add question number and content
            para = doc.add_paragraph()
            para.add_run(f"Question {question_data['number']}").bold = True
            
            doc.add_paragraph(question_data['content'].strip())
            
            # Add options without indicating correct answer
            for option in question_data['options']:
                doc.add_paragraph(option.strip())
                
            # Add spacing between questions
            doc.add_paragraph()
            
        doc.save(output_file)
        return True
        
    except Exception as e:
        print(f"Error creating practice document: {str(e)}")
        raise

def create_answer_key_document(input_file, output_file):
    """Creates a document with answers marked"""
    try:
        doc = Document()
        
        # Read and parse input file
        with open(input_file, 'r', encoding='utf-8', errors='replace') as file:
            content = file.read()
            
        # Split into questions
        questions = content.split('###################################################################')
        
        for question_text in questions:
            if not question_text.strip():
                continue
                
            question_data = parse_question(question_text)
            if not question_data:
                continue
                
            # Add question number and content
            para = doc.add_paragraph()
            para.add_run(f"Question {question_data['number']}").bold = True
            
            doc.add_paragraph(question_data['content'].strip())
            
            # Add options
            for option in question_data['options']:
                doc.add_paragraph(option.strip())
            
            # Add a line indicating correct answers
            correct_answers = ', '.join(question_data['correct'])
            answer_line = f"Correct Answers: {correct_answers}."
            
            # Check for most voted options
            most_voted = [opt for opt in question_data['options'] if 'Most Voted' in opt]
            if most_voted:
                most_voted_answers = ', '.join([opt[0] for opt in most_voted])  # Get the option letters
                answer_line += f" Most Voted: {most_voted_answers}."
            
            doc.add_paragraph(answer_line)
            
            # Add spacing between questions
            doc.add_paragraph()
            
        doc.save(output_file)
        return True
        
    except Exception as e:
        print(f"Error creating answer key document: {str(e)}")
        raise

def set_paragraph_format(paragraph):
    """Set the font and spacing for a paragraph."""
    # Set font to Arial, size 12
    run = paragraph.add_run()
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    
    # Set line spacing to 1.15
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = Pt(15)  # 1.15 line spacing
    paragraph_format.space_after = Pt(0)  # No space after
    paragraph_format.space_before = Pt(0)  # No space before

def create_snowpro_core_with_answers(input_file, output_file):
    """Creates a document with answers marked for Snowpro Core."""
    try:
        doc = Document()
        
        # Read and parse input file
        with open(input_file, 'r', encoding='utf-8', errors='replace') as file:
            content = file.read()
            
        # Split into questions
        questions = content.split('###################################################################')
        
        for question_text in questions:
            if not question_text.strip():
                continue
                
            question_data = parse_question(question_text)
            if not question_data:
                continue
                
            # Add question number and content
            para = doc.add_paragraph()
            set_paragraph_format(para)
            para.add_run(f"Question {question_data['number']}").bold = True
            
            para = doc.add_paragraph(question_data['content'].strip())
            set_paragraph_format(para)
            
            # Add options
            for option in question_data['options']:
                para = doc.add_paragraph(option.strip())
                set_paragraph_format(para)
            
            # Add a line indicating correct answers
            correct_answers = ', '.join(question_data['correct'])
            answer_line = f"Correct Answers: {correct_answers}."
            
            # Check for most voted options
            most_voted = [opt for opt in question_data['options'] if 'Most Voted' in opt]
            if most_voted:
                most_voted_answers = ', '.join([opt[0] for opt in most_voted])  # Get the option letters
                answer_line += f" Most Voted: {most_voted_answers}."
            
            para = doc.add_paragraph(answer_line)
            set_paragraph_format(para)
            
            # Add spacing between questions
            doc.add_paragraph()  # This adds a blank paragraph for spacing
            
        doc.save(output_file)
        return True
        
    except Exception as e:
        print(f"Error creating Snowpro Core document: {str(e)}")
        raise
